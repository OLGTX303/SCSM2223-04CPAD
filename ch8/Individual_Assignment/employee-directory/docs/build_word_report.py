from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "Chapter8_A24CS4045_ZUO_BOYU_Report.docx"
SCREENSHOT = ROOT / "docs" / "ui-screenshot.png"
TABLET = ROOT / "docs" / "tablet-screenshot.png"


def set_font(run, size=None, bold=False, color=None):
    run.font.name = "Calibri"
    if size:
        run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*color)


def set_paragraph(paragraph, before=0, after=6, line_spacing=1.1):
    fmt = paragraph.paragraph_format
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)
    fmt.line_spacing = line_spacing


def add_heading(doc, text, level=1):
    paragraph = doc.add_paragraph()
    set_paragraph(paragraph, before=10 if level == 1 else 6, after=4)
    run = paragraph.add_run(text)
    set_font(run, size=15 if level == 1 else 12, bold=True, color=(46, 116, 181))
    return paragraph


def add_body(doc, text, bold_label=None):
    paragraph = doc.add_paragraph()
    set_paragraph(paragraph)
    if bold_label:
        label = paragraph.add_run(bold_label)
        set_font(label, size=10.5, bold=True)
    run = paragraph.add_run(text)
    set_font(run, size=10.5)
    return paragraph


def add_bullet(doc, text, bold_label=None):
    paragraph = doc.add_paragraph(style="List Bullet")
    set_paragraph(paragraph, after=3)
    if bold_label:
        label = paragraph.add_run(bold_label)
        set_font(label, size=10.3, bold=True)
    run = paragraph.add_run(text)
    set_font(run, size=10.3)
    return paragraph


def shade_cell(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def style_table(table):
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    for row_index, row in enumerate(table.rows):
        for cell in row.cells:
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            for paragraph in cell.paragraphs:
                set_paragraph(paragraph, after=0)
                for run in paragraph.runs:
                    set_font(run, size=9, bold=row_index == 0)
            if row_index == 0:
                shade_cell(cell, "F2F4F7")


doc = Document()
section = doc.sections[0]
section.top_margin = Inches(0.7)
section.bottom_margin = Inches(0.65)
section.left_margin = Inches(0.75)
section.right_margin = Inches(0.75)

styles = doc.styles
styles["Normal"].font.name = "Calibri"
styles["Normal"].font.size = Pt(10.5)

title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
set_paragraph(title, after=2)
run = title.add_run("Chapter 8 Individual Assignment Report")
set_font(run, size=18, bold=True, color=(15, 23, 42))

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
set_paragraph(subtitle, after=8)
run = subtitle.add_run("Employee Directory with Vue 3, Axios, Express and MySQL")
set_font(run, size=11, color=(85, 95, 110))

meta = doc.add_table(rows=2, cols=4)
meta.columns[0].width = Inches(1.2)
meta.columns[1].width = Inches(2.1)
meta.columns[2].width = Inches(1.2)
meta.columns[3].width = Inches(2.0)
meta.cell(0, 0).text = "Name"
meta.cell(0, 1).text = "ZUO BOYU"
meta.cell(0, 2).text = "Matric No."
meta.cell(0, 3).text = "A24CS4045"
meta.cell(1, 0).text = "Course"
meta.cell(1, 1).text = "Cross Platform Application Development"
meta.cell(1, 2).text = "Chapter"
meta.cell(1, 3).text = "8 - Connecting Vue to Backend"
style_table(meta)

add_heading(doc, "1. Working UI Screenshots", 1)
add_body(
    doc,
    "Figure 1 shows the laptop view with summary cards, the add/edit form, search and sort controls, the paginated table, active/inactive badges, and Malaysian Ringgit salary formatting. Figure 2 shows the same UI re-flowed at tablet width: the form stacks above the table and no horizontal scrolling is required."
)
if SCREENSHOT.exists():
    picture_paragraph = doc.add_paragraph()
    picture_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    picture_paragraph.add_run().add_picture(str(SCREENSHOT), width=Inches(6.3))
    caption = doc.add_paragraph()
    caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph(caption, after=4)
    run = caption.add_run("Figure 1. Laptop view (>=1024 px): summary cards, form, paginated table, status badges, RM salaries.")
    set_font(run, size=8.8, color=(85, 95, 110))

if TABLET.exists():
    picture_paragraph = doc.add_paragraph()
    picture_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    picture_paragraph.add_run().add_picture(str(TABLET), width=Inches(3.6))
    caption = doc.add_paragraph()
    caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph(caption, after=2)
    run = caption.add_run("Figure 2. Tablet view (>=768 px): the workspace grid collapses to a single column.")
    set_font(run, size=8.8, color=(85, 95, 110))

doc.add_page_break()

add_heading(doc, "2. Chapter 8 Learning Outcomes - Code Evidence", 1)
table = doc.add_table(rows=1, cols=3)
table.columns[0].width = Inches(1.3)
table.columns[1].width = Inches(2.7)
table.columns[2].width = Inches(2.5)
hdr = table.rows[0].cells
hdr[0].text = "Learning Outcome"
hdr[1].text = "How the project satisfies it"
hdr[2].text = "Where to look (file:line)"
rows = [
    (
        "LO1. Connect Vue to a REST backend.",
        "App.vue owns employees and loading flags; onMounted triggers loadEmployees(); errors surface in a banner.",
        "src/App.vue:28, 30-53, 148",
    ),
    (
        "LO2. Use Axios with interceptors and async/await.",
        "Single axios.create() instance with baseURL/timeout/headers. Request interceptor logs every call; response interceptor maps every error shape to {message, status, errors}. All call sites are async/await.",
        "src/services/api.js:3, 12, 18, 35-52",
    ),
    (
        "LO3. Validated Vue forms with v-model and emits.",
        "EmployeeForm.vue uses v-model.trim on text inputs and v-model.number on Salary. validate() enforces all seven rules; inline field errors render before submit; payload is emitted to the parent.",
        "src/components/EmployeeForm.vue:73-94, 135, 177, 50",
    ),
    (
        "LO4. CRUD against MySQL with prepared statements, search and sort.",
        "Express implements GET/POST/PUT/DELETE /employees. LIKE search on 4 columns; ORDER BY restricted to a whitelist; all values use ? placeholders, including LIMIT/OFFSET; mysql2/promise pool.",
        "server/index.js:12, 26-73, 89, 120, 154; server/db.js:4",
    ),
]
for values in rows:
    cells = table.add_row().cells
    for index, value in enumerate(values):
        cells[index].text = value
style_table(table)

add_heading(doc, "3. Architecture and Component Structure", 1)
add_bullet(doc, "App.vue owns employees, filters, pagination metadata, and the currently selected employee. Children never call the API directly.", bold_label="Parent state: ")
add_bullet(doc, "EmployeeForm emits save/cancel; EmployeeList emits edit/delete/page-change; SearchSortControls emits change with debounced search and whitelisted sort keys.", bold_label="Children and emits: ")
add_bullet(doc, "src/services/api.js is the only file that imports Axios. Components import named functions and stay transport-agnostic.", bold_label="Service layer: ")
add_bullet(doc, "Validation runs twice: in EmployeeForm.validate() for inline feedback, and in server validateEmployee() as the authoritative gate before any SQL write. Duplicate empId/email becomes a friendly per-field error via the ER_DUP_ENTRY handler.", bold_label="Validation boundary: ")

doc.add_page_break()

add_heading(doc, "4. Challenges and Resolutions", 1)
add_bullet(doc, "Mapped req.query.sortBy through a fixed whitelist object; unknown keys fall back to name. The direction is reduced to the literal 'ASC' or 'DESC' before interpolation.", bold_label="Safe sorting: ")
add_bullet(doc, "Server validation errors, 5xx responses, DNS failures, and timeouts each produce a different Axios error shape. The response interceptor normalises them into one {message, status, errors} object the UI consumes.", bold_label="Heterogeneous Axios errors: ")
add_bullet(doc, "Node 18+ resolves localhost to ::1 while Laragon MySQL binds 127.0.0.1. The pool and CORS allow-list pin to 127.0.0.1 to avoid intermittent ECONNREFUSED.", bold_label="IPv6 vs IPv4 on Windows/Laragon: ")
add_bullet(doc, "The workspace grid collapses to a single column at <=768 px and the table sits in a horizontally-safe container, keeping the Actions column reachable without overflow.", bold_label="Responsive table on tablet: ")

add_heading(doc, "5. Extensions Beyond the Brief", 1)
add_bullet(doc, "Server-side pagination with clamped integers and a single count query that returns total/active/inactive in one round trip.")
add_bullet(doc, "Search extended to a fourth column (department) on top of the required three.")
add_bullet(doc, "Defensive server-side re-validation that mirrors the client rules, keeping the API safe even when called outside the Vue app.")
add_bullet(doc, "Summary cards driven by SUM(CASE WHEN active...) so totals stay consistent with the active filter.")

add_heading(doc, "6. Sources and References", 1)
add_body(doc, "Official documentation consulted during implementation. All code was written by me; AI assistants were used only for explanations and debugging.")
add_bullet(doc, "Vue 3 - Composition API and <script setup>: https://vuejs.org/guide/")
add_bullet(doc, "Vite configuration: https://vitejs.dev/config/")
add_bullet(doc, "Axios - instance configuration and interceptors: https://axios-http.com/docs/interceptors")
add_bullet(doc, "Express 5 routing and middleware: https://expressjs.com/en/5x/api.html")
add_bullet(doc, "mysql2 promise pool and prepared statements: https://github.com/sidorares/node-mysql2#using-prepared-statements")
add_bullet(doc, "MDN Intl.NumberFormat for MYR formatting: https://developer.mozilla.org/en-US/docs/Web/JavaScript/Reference/Global_Objects/Intl/NumberFormat")

doc.save(OUT)
print(OUT)
